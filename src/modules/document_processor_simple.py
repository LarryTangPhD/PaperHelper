#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简化文档处理模块
支持PDF、Word文档的读取、解析和处理
"""

import os
import io
import base64
from typing import Dict, List, Tuple, Optional, Any
import streamlit as st

# PDF处理
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word文档处理
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 简化文本处理（不使用spacy和nltk）
TEXT_PROCESSING_AVAILABLE = True

class SimpleDocumentProcessor:
    """简化文档处理器类"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.supported_formats = {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'txt': True,
            'image': False  # 暂时禁用OCR功能
        }
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """获取支持的文档格式"""
        return self.supported_formats
    
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """处理上传的文件"""
        if uploaded_file is None:
            return {"error": "没有上传文件"}
        
        file_info = {
            "filename": uploaded_file.name,
            "file_type": uploaded_file.type,
            "file_size": uploaded_file.size
        }
        
        try:
            # 根据文件类型处理
            if uploaded_file.type == "application/pdf":
                return self._process_pdf(uploaded_file, file_info)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._process_docx(uploaded_file, file_info)
            elif uploaded_file.type == "text/plain":
                return self._process_txt(uploaded_file, file_info)
            else:
                return {"error": f"不支持的文件格式: {uploaded_file.type}"}
        
        except Exception as e:
            return {"error": f"处理文件时出错: {str(e)}"}
    
    def _process_pdf(self, uploaded_file, file_info: Dict) -> Dict[str, Any]:
        """处理PDF文件"""
        if not PDF_AVAILABLE:
            return {"error": "PDF处理功能未安装"}
        
        try:
            # 读取PDF内容
            pdf_content = ""
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            
            # 提取文本
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                pdf_content += f"\n--- 第{page_num + 1}页 ---\n{page_text}"
            
            # 使用pdfplumber进行更详细的提取
            uploaded_file.seek(0)  # 重置文件指针
            with pdfplumber.open(uploaded_file) as pdf:
                detailed_content = ""
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        detailed_content += f"\n--- 第{page_num + 1}页 ---\n{page_text}"
            
            # 分析文档结构
            structure = self._analyze_document_structure(detailed_content)
            
            return {
                "success": True,
                "file_info": file_info,
                "content": detailed_content,
                "raw_content": pdf_content,
                "structure": structure,
                "word_count": len(detailed_content.split()),
                "page_count": len(pdf_reader.pages)
            }
        
        except Exception as e:
            return {"error": f"PDF处理失败: {str(e)}"}
    
    def _process_docx(self, uploaded_file, file_info: Dict) -> Dict[str, Any]:
        """处理Word文档"""
        if not DOCX_AVAILABLE:
            return {"error": "Word文档处理功能未安装"}
        
        try:
            # 读取Word文档
            doc = Document(uploaded_file)
            
            # 提取文本内容
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 提取表格内容
            table_content = ""
            for table in doc.tables:
                table_content += "\n--- 表格 ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_content += row_text + "\n"
            
            full_content = content + table_content
            
            # 分析文档结构
            structure = self._analyze_document_structure(full_content)
            
            return {
                "success": True,
                "file_info": file_info,
                "content": full_content,
                "raw_content": content,
                "table_content": table_content,
                "structure": structure,
                "word_count": len(full_content.split()),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        
        except Exception as e:
            return {"error": f"Word文档处理失败: {str(e)}"}
    
    def _process_txt(self, uploaded_file, file_info: Dict) -> Dict[str, Any]:
        """处理文本文件"""
        try:
            # 读取文本内容
            content = uploaded_file.read().decode('utf-8')
            
            # 分析文档结构
            structure = self._analyze_document_structure(content)
            
            return {
                "success": True,
                "file_info": file_info,
                "content": content,
                "raw_content": content,
                "structure": structure,
                "word_count": len(content.split()),
                "line_count": len(content.split('\n'))
            }
        
        except Exception as e:
            return {"error": f"文本文件处理失败: {str(e)}"}
    
    def _analyze_document_structure(self, content: str) -> Dict[str, Any]:
        """分析文档结构"""
        if not content:
            return {}
        
        structure = {
            "title": "",
            "sections": [],
            "paragraphs": [],
            "keywords": [],
            "summary": ""
        }
        
        try:
            # 分割段落
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            structure["paragraphs"] = paragraphs
            
            # 提取标题（假设第一段是标题）
            if paragraphs:
                structure["title"] = paragraphs[0][:100]  # 限制标题长度
            
            # 识别章节
            sections = []
            current_section = {"title": "", "content": ""}
            
            for para in paragraphs:
                # 简单的章节识别逻辑
                if (para.startswith('第') and ('章' in para or '节' in para)) or \
                   (len(para) < 50 and para.endswith('：')):
                    if current_section["title"]:
                        sections.append(current_section)
                    current_section = {"title": para, "content": ""}
                else:
                    current_section["content"] += para + "\n"
            
            if current_section["title"]:
                sections.append(current_section)
            
            structure["sections"] = sections
            
            # 提取关键词（简单实现）
            words = content.split()
            word_freq = {}
            for word in words:
                if len(word) > 1:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # 获取最常见的词作为关键词
            keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
            structure["keywords"] = [word for word, freq in keywords]
            
            # 生成摘要
            if len(content) > 200:
                structure["summary"] = content[:200] + "..."
            else:
                structure["summary"] = content
            
        except Exception as e:
            st.warning(f"文档结构分析时出错: {str(e)}")
        
        return structure
    
    def extract_citations(self, content: str) -> List[Dict[str, str]]:
        """提取引用信息"""
        citations = []
        
        # 简单的引用识别模式
        import re
        
        # 匹配常见的引用格式
        patterns = [
            r'\(([^)]+)\)',  # (作者, 年份)
            r'\[([^\]]+)\]',  # [1], [2]
            r'（([^）]+)）',  # 中文括号
            r'【([^】]+)】',  # 中文方括号
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                citations.append({
                    "text": match,
                    "type": "citation",
                    "position": content.find(match)
                })
        
        return citations
    
    def analyze_writing_style(self, content: str) -> Dict[str, Any]:
        """分析写作风格"""
        if not content:
            return {}
        
        # 简单的句子分割
        sentences = content.replace('。', '.').replace('！', '!').replace('？', '?').split('.')
        sentences = [s.strip() for s in sentences if s.strip()]
        
        analysis = {
            "word_count": len(content.split()),
            "sentence_count": len(sentences),
            "avg_sentence_length": 0,
            "formal_degree": 0,
            "complexity_score": 0
        }
        
        try:
            # 计算平均句子长度
            if sentences:
                analysis["avg_sentence_length"] = sum(len(s.split()) for s in sentences) / len(sentences)
            
            # 简单的正式程度评估
            formal_words = ['因此', '然而', '此外', '综上所述', '研究表明', '根据', '由于']
            formal_count = sum(1 for word in formal_words if word in content)
            analysis["formal_degree"] = min(formal_count / max(len(content.split()), 1) * 100, 100)
            
            # 复杂度评分
            analysis["complexity_score"] = min(analysis["avg_sentence_length"] * 10, 100)
            
        except Exception as e:
            st.warning(f"写作风格分析时出错: {str(e)}")
        
        return analysis

# 创建全局实例
document_processor = SimpleDocumentProcessor()
